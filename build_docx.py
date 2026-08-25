from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Cm(1.8)
sec.bottom_margin = Cm(1.8)
sec.left_margin   = Cm(2.0)
sec.right_margin  = Cm(2.0)

# ── Colour helpers ────────────────────────────────────────────────────────────
def rgb(h): return RGBColor((h>>16)&0xff,(h>>8)&0xff,h&0xff)

VIOLET  = rgb(0x7c3aed)
GLOW    = rgb(0xa78bfa)
DARK    = rgb(0x1e1b4b)
WHITE   = rgb(0xffffff)
BLACK   = rgb(0x0f172a)
SLATE   = rgb(0x334155)
GREEN   = rgb(0x059669)
AMBER   = rgb(0xd97706)
RED     = rgb(0xdc2626)
GREY    = rgb(0x64748b)
LBLUE   = rgb(0xddd6fe)   # light violet for header bg

# ── Para style helper ─────────────────────────────────────────────────────────
def para(text_val, size=11, bold=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text_val)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name  = "Calibri"
    return p

def heading(text_val, level=1):
    sizes  = {1: 22, 2: 14, 3: 12}
    colors = {1: VIOLET, 2: VIOLET, 3: SLATE}
    space_b = {1: 14, 2: 10, 3: 6}
    p = para(text_val, size=sizes[level], bold=True,
             color=colors[level], space_before=space_b[level], space_after=4)
    if level == 1:
        # Underline via border on paragraph
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'),  '6')
        bottom.set(qn('w:space'),'1')
        bottom.set(qn('w:color'),'7c3aed')
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])

def cell_text(cell, text_val, size=10, bold=False, color=BLACK,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text_val)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic= italic
    run.font.color.rgb = color
    run.font.name  = "Calibri"

def make_table(headers, rows, col_widths, header_color="3730a3",
               header_text_color=WHITE):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shading(cell, header_color)
        cell_text(cell, h, size=10, bold=True, color=header_text_color,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri+1]
        bg = "f5f3ff" if ri % 2 == 0 else "ffffff"
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            shading(cell, bg)
            # Pick color for status column
            fc = BLACK
            if isinstance(val, tuple):
                val, fc = val
            cell_text(cell, val, size=10, color=fc)

    set_col_widths(table, col_widths)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
para("HADES — AI OPERATING SYSTEM", size=28, bold=True, color=VIOLET,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
para("Complete Technology Stack", size=14, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
para("Linux-Native  ·  Model-Agnostic  ·  Mission-Locked  ·  Verified Execution",
     size=11, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=0, space_after=12)


# ══════════════════════════════════════════════════════════════════════════════
# 1. BACKEND — PYTHON
# ══════════════════════════════════════════════════════════════════════════════
heading("1.  Backend — Python", 1)

make_table(
    headers=["Package", "Version", "Role"],
    rows=[
        ("FastAPI",       "0.141.1", "Web framework — REST API endpoints + SSE streaming"),
        ("Uvicorn",       "0.52.2",  "ASGI server — runs the FastAPI application"),
        ("Pydantic",      "2.13.4",  "Data models — Mission, MissionUnderstanding, SessionState"),
        ("LiteLLM",       "1.96.2",  "Unified LLM abstraction — wraps all model providers with auto-fallback"),
        ("python-dotenv", "1.2.2",   "Loads API keys from .env file into environment"),
        ("asyncio",       "stdlib",  "Background task execution, subprocess shell management"),
        ("pytest",        "latest",  "Testing framework (1 test file currently: test_partner_brain.py)"),
    ],
    col_widths=[4.5, 2.5, 10.0],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 2. AI / LLM PROVIDERS
# ══════════════════════════════════════════════════════════════════════════════
heading("2.  AI / LLM Providers  (via LiteLLM)", 1)

make_table(
    headers=["Provider", "Model", "Capability", "Status"],
    rows=[
        ("Google AI Studio", "gemini/gemini-2.5-flash-lite",    "Conversational / Partner Brain",    "Active"),
        ("Google AI Studio", "gemini/gemini-flash-latest",       "Research / Analysis tasks",         "Active"),
        ("Groq",             "llama-3.3-70b-versatile",          "Fast open-source inference",        "Active"),
        ("OpenRouter",       "llama-3.2-3b-instruct",            "Free public endpoint",              "Active"),
        ("Ollama (local)",   "user-configured local model",      "100% offline / private execution",  "Partial"),
        ("OpenAI",           "gpt-4o (optional)",                "Premium reasoning (optional)",      "Config only"),
        ("Anthropic",        "claude (optional)",                "Premium reasoning (optional)",      "Config only"),
    ],
    col_widths=[3.5, 5.0, 5.5, 3.0],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 3. FRONTEND
# ══════════════════════════════════════════════════════════════════════════════
heading("3.  Frontend — React / TypeScript / Vite", 1)

heading("3a.  Core Framework", 2)
make_table(
    headers=["Package", "Version", "Role"],
    rows=[
        ("React",          "18.3.1",  "UI component framework"),
        ("React DOM",      "18.3.1",  "DOM rendering for React"),
        ("TypeScript",     "5.5.4",   "Type-safe JavaScript for all frontend code"),
        ("Vite",           "5.2.0",   "Build tool and hot-reload dev server"),
        ("@vitejs/plugin-react", "4.2.1", "Vite plugin for React/JSX support"),
    ],
    col_widths=[5.0, 2.5, 9.5],
)

para("")
heading("3b.  Styling", 2)
make_table(
    headers=["Package", "Version", "Role"],
    rows=[
        ("Tailwind CSS",  "3.4.17",  "Utility-first CSS framework — all UI styling"),
        ("PostCSS",       "latest",  "CSS transformation pipeline"),
        ("Autoprefixer",  "latest",  "Adds vendor prefixes for CSS compatibility"),
        ("@emotion/react","11.13.3", "CSS-in-JS for dynamic styles"),
    ],
    col_widths=[5.0, 2.5, 9.5],
)

para("")
heading("3c.  UI Libraries", 2)
make_table(
    headers=["Package", "Version", "Role"],
    rows=[
        ("Lucide React",   "0.522.0", "Icon library — all UI icons"),
        ("Framer Motion",  "11.5.4",  "Animation library — UI transitions and motion"),
    ],
    col_widths=[5.0, 2.5, 9.5],
)

para("")
heading("3d.  Dev Tooling", 2)
make_table(
    headers=["Package", "Version", "Role"],
    rows=[
        ("ESLint",                       "8.50.0", "JavaScript/TypeScript linter"),
        ("@typescript-eslint/parser",    "5.54.0", "TypeScript-aware ESLint parser"),
        ("eslint-plugin-react-hooks",    "4.6.0",  "Lint rules for React hooks"),
        ("eslint-plugin-react-refresh",  "0.4.1",  "Lint rules for Vite React Refresh"),
        ("@types/react",                 "18.3.1", "TypeScript type definitions for React"),
        ("@types/node",                  "20.11.18","TypeScript type definitions for Node.js"),
    ],
    col_widths=[5.5, 2.5, 9.0],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 4. VOICE / AUDIO
# ══════════════════════════════════════════════════════════════════════════════
heading("4.  Voice & Audio", 1)

make_table(
    headers=["Component", "Technology", "Role", "Runs On"],
    rows=[
        ("Text-to-Speech (TTS)", "Kokoro-ONNX",         "Local neural TTS — synthesizes audio on backend server", "Linux backend"),
        ("Audio Delivery",       "Base64 + JSON",        "TTS audio encoded as Base64, sent in HTTP response",     "API response"),
        ("Audio Playback",       "HTML5 Audio API",      "Browser plays Base64 audio received from backend",       "Browser"),
        ("Speech-to-Text (STT)", "Web Speech API",       "Browser-native voice input — window.SpeechRecognition",  "Browser (Chrome)"),
        ("Backend STT",          "Whisper (planned)",    "Server-side offline STT — FUTURE MILESTONE",             "Not yet"),
    ],
    col_widths=[4.0, 3.5, 6.5, 3.0],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 5. COMMUNICATION PROTOCOLS
# ══════════════════════════════════════════════════════════════════════════════
heading("5.  Communication Protocols", 1)

make_table(
    headers=["Protocol", "Endpoint", "Direction", "Purpose"],
    rows=[
        ("HTTP REST (POST)",       "/api/chat",        "Browser → Backend", "Send user messages, receive AI response"),
        ("Server-Sent Events (SSE)", "/api/events",    "Backend → Browser", "Real-time mission execution updates"),
        ("HTTP GET",               "/api/config/status","Browser → Backend", "Check which API keys are configured"),
        ("localStorage",           "(browser storage)","Client-side",       "Persist session ID and user name"),
    ],
    col_widths=[4.5, 3.5, 3.5, 5.5],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 6. LINUX INTEGRATION
# ══════════════════════════════════════════════════════════════════════════════
heading("6.  Linux System Integration", 1)

make_table(
    headers=["Tool / API", "Python Module", "Role", "File"],
    rows=[
        ("Bash shell execution",   "asyncio.create_subprocess_shell", "Run any bash command on the host Linux system",     "terminal.py"),
        ("POSIX filesystem",       "os, pathlib (stdlib)",            "Read, write, move, inspect local files",            "filesystem.py"),
        ("Process management",     "psutil / subprocess (stdlib)",    "List running processes, monitor system resources",   "process_manager.py"),
        ("Local LLM server",       "HTTP to localhost:11434",         "Ollama serves AI models 100% offline on Linux",     "worker_manager.py"),
        ("Environment variables",  "python-dotenv + os.environ",      "Load API keys from .env without hardcoding secrets", "main.py"),
    ],
    col_widths=[4.5, 4.5, 6.0, 2.0],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 7. CONFIGURATION & DATA
# ══════════════════════════════════════════════════════════════════════════════
heading("7.  Configuration & Data Files", 1)

make_table(
    headers=["File", "Format", "Role"],
    rows=[
        ("config.json",      "JSON",      "Worker definitions — model names, capability tags, API key env var mappings, fallback order"),
        (".env",             "Key=Value", "Secret API keys — GEMINI_API_KEY, GROQ_API_KEY, OPENROUTER_API_KEY, OLLAMA_API_BASE"),
        ("memory.json",      "JSON",      "Flat-file mission history — mission summaries, outputs, outcomes (appended on completion)"),
        ("requirements.txt", "Plain text","Python backend dependencies"),
        ("package.json",     "JSON",      "Node.js frontend dependencies and Vite build scripts"),
        ("tailwind.config.js","JS",       "Tailwind CSS theme — custom color tokens (void, ion, signal, glow, etc.)"),
        ("vite.config.ts",   "TypeScript","Vite build configuration and React plugin setup"),
    ],
    col_widths=[4.0, 2.5, 10.5],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 8. CORE BACKEND MODULES
# ══════════════════════════════════════════════════════════════════════════════
heading("8.  Core Backend Modules (Custom Code)", 1)

make_table(
    headers=["Module", "File", "Role"],
    rows=[
        ("Partner Brain",             "src/core/partner_brain.py",              "Conversational AI coordinator — routes intent, manages alignment"),
        ("Intent Classifier",         "src/core/intent_classifier.py",          "LLM-based: casual / small task / real mission classification"),
        ("Mission Extractor",         "src/core/mission_extractor.py",          "Populates MissionUnderstanding fields from conversation"),
        ("Understanding Evaluator",   "src/core/understanding_evaluator.py",    "Checks if all mission fields populated → triggers Mission Lock"),
        ("Conversational Decision",   "src/core/conversational_decision.py",    "Generates ASK / PROPOSE / ACKNOWLEDGE action from system prompt"),
        ("Execution Brain",           "src/core/execution/execution_brain.py",  "Async task orchestration — TaskGraph, worker delegation, retry"),
        ("Worker Manager",            "src/core/worker_manager.py",             "LiteLLM wrapper — capability routing, provider fallback"),
        ("Review Engine",             "src/core/execution/review_engine.py",    "4-stage output validation before mission delivery"),
        ("Memory Manager",            "src/core/memory_manager.py",             "Session state + JSON mission history read/write"),
        ("Voice Manager",             "src/core/voice_manager.py",              "Kokoro-ONNX TTS — generates Base64 audio on the backend"),
        ("Terminal Skill",            "src/skills/computer/terminal.py",        "Executes bash commands via asyncio subprocess on Linux"),
        ("Filesystem Skill",          "src/skills/computer/filesystem.py",      "Reads, writes, lists files on the local Linux filesystem"),
        ("Process Manager Skill",     "src/skills/computer/process_manager.py", "Lists and monitors running Linux processes"),
        ("Skill Registry",            "src/skills/registry.py",                 "Maps capability strings to skill adapter classes"),
        ("Mission Model",             "src/models/mission.py",                  "Pydantic schema — Mission, MissionStatus enum, MissionUnderstanding"),
        ("Conversation Model",        "src/models/conversation.py",             "Pydantic schema — ConversationMessage, SessionState"),
        ("FastAPI Entry Point",        "main.py",                               "API routes /api/chat, /api/events, /api/config/status"),
    ],
    col_widths=[4.5, 5.5, 7.0],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 9. CORE FRONTEND MODULES
# ══════════════════════════════════════════════════════════════════════════════
heading("9.  Core Frontend Modules (Custom Code)", 1)

make_table(
    headers=["Module", "File", "Role"],
    rows=[
        ("HadesService",    "frontend/src/services/HadesService.ts",  "Singleton service — all API calls, SSE, audio, voice recognition, state"),
        ("useHades hook",   "frontend/src/services/useHades.ts",      "React hook — subscribes to HadesService state updates"),
        ("App",             "frontend/src/App.tsx",                   "Root component — name init flow + main layout"),
        ("ChatPanel",       "frontend/src/components/ChatPanel.tsx",  "Message history display — user and Hades messages"),
        ("Composer",        "frontend/src/components/Composer.tsx",   "Input box — text + send + mic + attach buttons"),
        ("HeroCore",        "frontend/src/components/HeroCore.tsx",   "Animated HADES central visual element"),
        ("HadesPresence",   "frontend/src/components/HadesPresence.tsx","Status indicator — idle / processing / executing"),
        ("TopBar",          "frontend/src/components/TopBar.tsx",     "Top navigation bar with system controls"),
        ("AgendaCard",      "frontend/src/components/AgendaCard.tsx", "Mission agenda sidebar panel"),
    ],
    col_widths=[3.5, 5.5, 8.0],
)

para("")

# ══════════════════════════════════════════════════════════════════════════════
# 10. SUMMARY COUNT
# ══════════════════════════════════════════════════════════════════════════════
heading("10.  Summary", 1)

make_table(
    headers=["Category", "Count", "Key Technologies"],
    rows=[
        ("Python backend packages",   "7",  "FastAPI, Uvicorn, Pydantic, LiteLLM, python-dotenv, asyncio, pytest"),
        ("AI / LLM providers",        "5",  "Google Gemini, Groq, OpenRouter, Ollama, OpenAI (optional)"),
        ("Frontend packages (prod)",  "5",  "React, TypeScript, Vite, Tailwind CSS, Lucide React, Framer Motion"),
        ("Frontend packages (dev)",   "6",  "ESLint, typescript-eslint, @types/react, @types/node, etc."),
        ("Voice / audio components",  "4",  "Kokoro-ONNX, Base64 transport, HTML5 Audio, Web Speech API"),
        ("Communication protocols",   "3",  "HTTP REST, Server-Sent Events, localStorage"),
        ("Linux integrations",        "5",  "asyncio subprocess, POSIX fs, psutil, Ollama, python-dotenv"),
        ("Custom backend modules",    "17", "Partner Brain, ExecutionBrain, WorkerManager, ReviewEngine, etc."),
        ("Custom frontend modules",   "9",  "HadesService, ChatPanel, Composer, HeroCore, AgendaCard, etc."),
        ("Config / data files",       "7",  "config.json, .env, memory.json, requirements.txt, package.json, etc."),
        ("TOTAL COMPONENTS",          "62+","Full-stack AI OS prototype"),
    ],
    col_widths=[5.5, 2.0, 9.5],
    header_color="4c1d95",
)

# Footer
para("")
para("HADES — AI Operating System Prototype  ·  Linux-Native  ·  2025",
     size=9, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=8, space_after=0)


# ── SAVE ──────────────────────────────────────────────────────────────────────
out = r"d:\HACK O HADES\HADES_TechStack.docx"
doc.save(out)
print(f"Saved: {out}")
